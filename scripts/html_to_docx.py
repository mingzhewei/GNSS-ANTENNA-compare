# -*- coding: utf-8 -*-
"""将 M21_双天线抗干扰对比测试_采集与判定方案.html 转换为 Word .docx。

用法（项目根目录）：
    .venv/Scripts/python.exe scripts/html_to_docx.py

要点：
- 显式以 UTF-8 读取源 HTML，杜绝乱码。
- 全部 run 同时设置 ascii 与 eastAsia 字体（微软雅黑），代码段用 Consolas+雅黑。
- 表格转 Word 原生 Table Grid 表格（含 rowspan 纵向合并），A4 页面。
"""
import os
import re
import sys

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# 默认转换 M21 方案；也可用命令行参数指定：html_to_docx.py <输入.html> <输出.docx>
DEFAULT_SRC = os.path.join(ROOT, "M21_双天线抗干扰对比测试_采集与判定方案.html")
DEFAULT_DST = os.path.join(ROOT, "M21_双天线抗干扰对比测试_采集与判定方案.docx")

FONT_CN = "Microsoft YaHei"   # 全文中文主字体
FONT_MONO = "Consolas"        # 代码/命令段西文字体
BODY_SIZE = Pt(10.5)          # 正文五号
TABLE_SIZE_DEFAULT = Pt(9)
TABLE_SIZE_WIDE = Pt(8)       # >=5 列的宽表
TABLE_SIZE_XWIDE = Pt(7.5)    # >=5 列且内容很长的表（第 1 节 Log 表）
FOOTER_SIZE = Pt(9)


def set_run_font(run, size=BODY_SIZE, bold=False, mono=False, color=None):
    """统一设置 run 字体：ascii/hAnsi 与 eastAsia 必须同时设置。"""
    run.font.size = size
    run.font.bold = bold
    name = FONT_MONO if mono else FONT_CN
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), FONT_CN)  # 中文一律雅黑
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, node, size=BODY_SIZE, bold=False, mono=False):
    """递归把 HTML 内联节点写入段落，保留 b/strong、code、sub/sup、a。"""
    for child in node.children:
        emit_inline(paragraph, child, size=size, bold=bold, mono=mono)


def emit_inline(paragraph, child, size=BODY_SIZE, bold=False, mono=False):
    """把单个节点（文本或内联标签）写入段落。"""
    if isinstance(child, NavigableString):
        text = str(child)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, mono=mono)
    elif isinstance(child, Tag):
        name = child.name
        if name == "br":
            paragraph.add_run().add_break()
        elif name in ("b", "strong"):
            add_inline(paragraph, child, size=size, bold=True, mono=mono)
        elif name == "code":
            add_inline(paragraph, child, size=size, bold=bold, mono=True)
        elif name == "sub":
            before = len(paragraph.runs)
            add_inline(paragraph, child, size=size, bold=bold, mono=mono)
            for run in paragraph.runs[before:]:
                run.font.subscript = True
        elif name == "sup":
            before = len(paragraph.runs)
            add_inline(paragraph, child, size=size, bold=bold, mono=mono)
            for run in paragraph.runs[before:]:
                run.font.superscript = True
        elif name == "a":
            add_inline(paragraph, child, size=size, bold=bold, mono=mono)
            href = child.get("href", "")
            if href.startswith("http"):
                run = paragraph.add_run("（" + href + "）")
                set_run_font(run, size=Pt(size.pt - 1.5), mono=True,
                             color=RGBColor(0x2B, 0x6C, 0xB0))
        else:  # span 等其余内联标签：只取文本
            add_inline(paragraph, child, size=size, bold=bold, mono=mono)


def add_para(doc, node, style=None, size=BODY_SIZE, bold=False,
             align=None, space_after=Pt(6)):
    p = doc.add_paragraph(style=style)
    add_inline(p, node, size=size, bold=bold)
    for run in p.runs:
        if run.font.size is None:
            set_run_font(run, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_heading(doc, node, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline(p, node, bold=True)
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    for run in p.runs:
        set_run_font(run, size=sizes[level], bold=True)
        run.font.color.rgb = RGBColor(0x1F, 0x23, 0x29)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def table_font_size(table_tag, ncols, grid):
    if ncols >= 5:
        longest = max((len(item[0].get_text()) for row in grid for item in row if item),
                      default=0)
        return TABLE_SIZE_XWIDE if longest > 150 else TABLE_SIZE_WIDE
    return TABLE_SIZE_DEFAULT


def build_grid(table_tag):
    """把 HTML 表格展开成网格，处理 rowspan/colspan。返回 (grid, ncols)。
    grid[r] = [(cell_tag, rowspan, colspan, is_origin), ...] 按列对齐
    （None 表示被上方 rowspan 占用）。"""
    rows = table_tag.find_all("tr")
    occupied = {}  # (row, col) -> True 被 rowspan 占用
    grid = []
    ncols = 0
    for r, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"], recursive=False)
        row_cells = []
        c = 0
        for tag in cells:
            while occupied.get((r, c)):
                row_cells.append(None)
                c += 1
            rowspan = int(tag.get("rowspan", 1))
            colspan = int(tag.get("colspan", 1))
            row_cells.append((tag, rowspan, colspan, True))
            for dr in range(1, rowspan):
                for dc in range(colspan):
                    occupied[(r + dr, c + dc)] = True
            for dc in range(1, colspan):
                row_cells.append((tag, rowspan, colspan, False))
                c += 1
            c += 1
        grid.append(row_cells)
        ncols = max(ncols, len(row_cells))
    for row in grid:
        row.extend([None] * (ncols - len(row)))
    return grid, ncols


def column_widths_cm(table_tag, ncols, total_cm=17.0):
    """根据首行 th/td 的 style width(px) 提示估算各列宽度（cm）。
    无提示的列平分剩余宽度；整体按 total_cm（A4 版心 17cm）归一。"""
    first_row = table_tag.find("tr")
    if not first_row:
        return None
    hints = []
    for cell in first_row.find_all(["th", "td"], recursive=False):
        style = cell.get("style") or ""
        m = re.search(r"width:\s*(\d+)px", style)
        hints.append(int(m.group(1)) if m else None)
    hints += [None] * (ncols - len(hints))
    if all(h is None for h in hints):
        return None
    spec = sum(h for h in hints if h)
    n_unspec = sum(1 for h in hints if h is None)
    # 未指定列按平均内容长度分配剩余 px（假设整表约 700px）
    fallback = max(80, (700 - spec) // max(n_unspec, 1))
    px = [h if h else fallback for h in hints]
    total = sum(px)
    return [total_cm * w / total for w in px]


def add_table(doc, table_tag):
    grid, ncols = build_grid(table_tag)
    nrows = len(grid)
    fsize = table_font_size(table_tag, ncols, grid)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    widths = column_widths_cm(table_tag, ncols)
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Cm(w)

    origin_cells = {}  # (r, c) -> (docx cell, rowspan, colspan)，供合并
    for r, row in enumerate(grid):
        for c, item in enumerate(row):
            if item is None:
                continue
            cell_tag, rowspan, colspan, is_origin = item
            if not is_origin:
                continue
            cell = table.cell(r, c)
            origin_cells[(r, c)] = (cell, rowspan, colspan)
            is_header = cell_tag.name == "th"
            p = cell.paragraphs[0]
            add_inline(p, cell_tag, size=fsize, bold=is_header)
            for run in p.runs:
                if run.font.size is None:
                    set_run_font(run, size=fsize, bold=is_header)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)

    # rowspan/colspan 合并
    for (r, c), (cell, rowspan, colspan) in origin_cells.items():
        if rowspan > 1 or colspan > 1:
            cell.merge(table.cell(r + rowspan - 1, c + colspan - 1))

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_cmd_block(doc, node):
    """div.cmd / div.tree 预格式命令块：逐行、等宽字体、小字号。"""
    for line in node.get_text().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        set_run_font(run, size=Pt(9), mono=True)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def add_dl(doc, node):
    """dl.kv：术语加粗 + 冒号 + 定义，一个 dt/dd 对一段。"""
    dts = node.find_all("dt", recursive=False)
    dds = node.find_all("dd", recursive=False)
    for dt, dd in zip(dts, dds):
        p = doc.add_paragraph()
        run = p.add_run(dt.get_text() + "：")
        set_run_font(run, bold=True)
        add_inline(p, dd)
        p.paragraph_format.space_after = Pt(4)


def add_list(doc, node, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        add_inline(p, li)
        for run in p.runs:
            if run.font.size is None:
                set_run_font(run)
        p.paragraph_format.space_after = Pt(3)


INLINE_TAGS = {"b", "strong", "span", "a", "code", "sub", "sup", "em", "i", "u"}


def walk(doc, node):
    """递归遍历块级元素。连续的内联内容（裸文本 + b/span/a 等）合并为一个段落。"""
    pending = []  # 待写入的内联节点缓冲

    def flush():
        if pending:
            p = doc.add_paragraph()
            for n in pending:
                emit_inline(p, n)
            p.paragraph_format.space_after = Pt(6)
            pending.clear()

    for child in node.children:
        if isinstance(child, NavigableString):
            if child.strip() or pending:
                pending.append(child)
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name
        classes = child.get("class", [])
        if name in INLINE_TAGS:
            pending.append(child)
            continue
        flush()
        if name == "h1":
            add_heading(doc, child, 1)
        elif name == "h2":
            add_heading(doc, child, 2)
        elif name == "h3":
            add_heading(doc, child, 3)
        elif name == "p":
            align = (WD_ALIGN_PARAGRAPH.CENTER
                     if "text-align:center" in (child.get("style") or "")
                     else None)
            size = FOOTER_SIZE if "sub" in classes else BODY_SIZE
            add_para(doc, child, size=size, align=align)
        elif name == "table":
            add_table(doc, child)
        elif name == "ul":
            add_list(doc, child, numbered=False)
        elif name == "ol":
            add_list(doc, child, numbered=True)
        elif name == "dl":
            add_dl(doc, child)
        elif name == "div" and ("cmd" in classes or "tree" in classes):
            add_cmd_block(doc, child)
        elif name in ("div", "header", "footer", "section", "article"):
            walk(doc, child)
        elif name == "nav":
            # 目录导航：转成一行文本
            p = doc.add_paragraph()
            run = p.add_run("目录：")
            set_run_font(run, bold=True)
            for a in child.find_all("a"):
                r = p.add_run(a.get_text() + "　")
                set_run_font(r, size=FOOTER_SIZE)
            p.paragraph_format.space_after = Pt(6)
        # script/style 等忽略
    flush()


def setup_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)  # A4
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # Normal 样式：正文 10.5pt，中文雅黑
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = BODY_SIZE
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_CN)
    rFonts.set(qn("w:hAnsi"), FONT_CN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    pf = normal.paragraph_format
    pf.line_spacing = 1.3


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_SRC
    dst = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_DST
    with open(src, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    doc = Document()
    setup_document(doc)
    walk(doc, soup.body)
    doc.save(dst)
    print(f"saved: {dst}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
